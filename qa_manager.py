import json
import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime

DATA_FILE = "qa_data_gui.json"


def load_data():
    if not os.path.exists(DATA_FILE):
        return {"test_scenarios": [], "test_cases": [], "bug_reports": []}
    with open(DATA_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def save_data(data):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


def generate_id(prefix, existing_items):
    number = len(existing_items) + 1
    return f"{prefix}{number:02d}"


class QAApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("QA Manager – Test Cases, Scenáre & Bugy")
        self.geometry("1000x650")

        self.data = load_data()
        self.dark_mode = False

        # výbery na úpravu / mazanie
        self.selected_ts_index = None
        self.selected_tc_id = None
        self.selected_bug_id = None
        self.bug_screenshot_path = None  # cesta k screenshotu pre bug

        self._styled_text_widgets = []
        self._styled_listbox_widgets = []

        self.create_menu()

        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True)

        self.create_ts_tab()
        self.create_tc_tab()
        self.create_bug_tab()

        self.apply_theme()

    # ===== MENU =====
    def create_menu(self):
        menubar = tk.Menu(self)

        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Export do TXT", command=self.export_to_txt)
        file_menu.add_command(label="Export do HTML", command=self.export_to_html)
        file_menu.add_command(label="Export do Word", command=self.export_to_word)
        file_menu.add_command(label="Export do PDF", command=self.export_to_pdf)
        file_menu.add_separator()
        file_menu.add_command(label="Resetovať databázu", command=self.reset_database)
        file_menu.add_separator()
        file_menu.add_command(label="Koniec", command=self.quit)
        menubar.add_cascade(label="Súbor", menu=file_menu)

        view_menu = tk.Menu(menubar, tearoff=0)
        view_menu.add_command(label="Prepnúť Dark Mode", command=self.toggle_dark_mode)
        menubar.add_cascade(label="Zobrazenie", menu=view_menu)

        self.config(menu=menubar)

    def toggle_dark_mode(self):
        self.dark_mode = not self.dark_mode
        self.apply_theme()

    def apply_theme(self):
        bg = "#1e1e1e" if self.dark_mode else "#f0f0f0"
        fg = "#ffffff" if self.dark_mode else "#000000"

        self.configure(bg=bg)

        def style_text_widget(widget):
            try:
                widget.configure(bg=bg, fg=fg, insertbackground=fg)
            except tk.TclError:
                pass

        def style_listbox(widget):
            try:
                widget.configure(bg=bg, fg=fg)
            except tk.TclError:
                pass

        for widget in self._styled_text_widgets:
            style_text_widget(widget)
        for widget in self._styled_listbox_widgets:
            style_listbox(widget)

    # ===== TAB – TEST SCENÁRE =====
    def create_ts_tab(self):
        self.ts_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.ts_tab, text="Test Scenáre (TS)")

        form_frame = ttk.LabelFrame(self.ts_tab, text="Pridať / Upraviť Test Scenár")
        form_frame.pack(side="top", fill="x", padx=10, pady=10)

        ttk.Label(form_frame, text="Názov TS:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.ts_title_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.ts_title_var, width=50).grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Popis:").grid(row=1, column=0, sticky="nw", padx=5, pady=5)
        self.ts_desc_text = tk.Text(form_frame, height=4, width=50)
        self.ts_desc_text.grid(row=1, column=1, padx=5, pady=5)

        ttk.Button(form_frame, text="Uložiť nový TS", command=self.add_ts).grid(
            row=2, column=1, sticky="e", padx=5, pady=5
        )
        ttk.Button(form_frame, text="Uložiť zmeny TS", command=self.update_ts).grid(
            row=3, column=1, sticky="e", padx=5, pady=5
        )

        list_frame = ttk.LabelFrame(self.ts_tab, text="Zoznam Test Scenárov")
        list_frame.pack(side="top", fill="both", expand=True, padx=10, pady=10)

        self.ts_list = tk.Listbox(list_frame)
        self.ts_list.pack(side="left", fill="both", expand=True, padx=(5, 0), pady=5)
        self.ts_list.bind("<<ListboxSelect>>", self.on_ts_select)

        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.ts_list.yview)
        scrollbar.pack(side="left", fill="y")
        self.ts_list.config(yscrollcommand=scrollbar.set)

        btn_frame = ttk.Frame(list_frame)
        btn_frame.pack(side="bottom", fill="x", padx=5, pady=5)
        ttk.Button(btn_frame, text="Vymazať vybraný TS", command=self.delete_ts).pack(side="right")

        self._styled_text_widgets.append(self.ts_desc_text)
        self._styled_listbox_widgets.append(self.ts_list)

        self.refresh_ts_list()

    def add_ts(self):
        title = self.ts_title_var.get().strip()
        desc = self.ts_desc_text.get("1.0", "end").strip()

        if not title:
            messagebox.showerror("Chyba", "Názov TS nemôže byť prázdny.")
            return

        ts_id = generate_id("TS", self.data["test_scenarios"])
        ts = {"id": ts_id, "title": title, "description": desc}
        self.data["test_scenarios"].append(ts)
        save_data(self.data)

        self.ts_title_var.set("")
        self.ts_desc_text.delete("1.0", "end")
        self.selected_ts_index = None

        self.refresh_ts_list()
        self.refresh_tc_ts_combobox()
        self.apply_theme()
        messagebox.showinfo("OK", f"Test scenár {ts_id} uložený.")

    def on_ts_select(self, event):
        selection = self.ts_list.curselection()
        if not selection:
            self.selected_ts_index = None
            return
        index = selection[0]
        if index >= len(self.data["test_scenarios"]):
            self.selected_ts_index = None
            return

        ts = self.data["test_scenarios"][index]
        self.selected_ts_index = index

        self.ts_title_var.set(ts["title"])
        self.ts_desc_text.delete("1.0", "end")
        self.ts_desc_text.insert("1.0", ts["description"])

    def update_ts(self):
        if self.selected_ts_index is None:
            messagebox.showerror("Chyba", "Najprv vyber test scenár zo zoznamu.")
            return

        title = self.ts_title_var.get().strip()
        desc = self.ts_desc_text.get("1.0", "end").strip()

        if not title:
            messagebox.showerror("Chyba", "Názov TS nemôže byť prázdny.")
            return

        ts = self.data["test_scenarios"][self.selected_ts_index]
        ts["title"] = title
        ts["description"] = desc
        save_data(self.data)
        self.refresh_ts_list()
        self.apply_theme()
        messagebox.showinfo("OK", f"Test scenár {ts['id']} bol upravený.")

    def refresh_ts_list(self):
        self.ts_list.delete(0, "end")
        for ts in self.data["test_scenarios"]:
            self.ts_list.insert("end", f"{ts['id']} – {ts['title']}")

    def delete_ts(self):
        selection = self.ts_list.curselection()
        if not selection:
            messagebox.showerror("Chyba", "Najprv vyber test scenár zo zoznamu.")
            return
        index = selection[0]
        if index >= len(self.data["test_scenarios"]):
            return
        ts = self.data["test_scenarios"][index]

        if not messagebox.askyesno(
            "Vymazať TS",
            f"Naozaj chceš vymazať {ts['id']} – {ts['title']}?\n\nTest Cases, ktoré naň odkazujú, zostanú nezmenené."
        ):
            return

        del self.data["test_scenarios"][index]
        save_data(self.data)
        self.refresh_ts_list()
        self.refresh_tc_ts_combobox()
        self.selected_ts_index = None
        messagebox.showinfo("OK", "Test scenár vymazaný.")

    # ===== TAB – TEST CASES =====
    def create_tc_tab(self):
        self.tc_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.tc_tab, text="Test Cases (TC)")

        top_frame = ttk.Frame(self.tc_tab)
        top_frame.pack(side="top", fill="x", padx=10, pady=10)

        form_frame = ttk.LabelFrame(top_frame, text="Pridať / Upraviť Test Case")
        form_frame.pack(side="left", fill="both", expand=True, padx=5, pady=5)

        ttk.Label(form_frame, text="Názov TC:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.tc_title_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.tc_title_var, width=40).grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Predpoklady:").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        self.tc_pre_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.tc_pre_var, width=40).grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="TS ID:").grid(row=2, column=0, sticky="w", padx=5, pady=5)
        self.tc_ts_var = tk.StringVar()
        self.tc_ts_combo = ttk.Combobox(form_frame, textvariable=self.tc_ts_var, width=37, state="readonly")
        self.tc_ts_combo.grid(row=2, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Kroky (každý na nový riadok):").grid(
            row=3, column=0, sticky="nw", padx=5, pady=5
        )
        self.tc_steps_text = tk.Text(form_frame, height=4, width=40)
        self.tc_steps_text.grid(row=3, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Očakávaný výsledok:").grid(row=4, column=0, sticky="w", padx=5, pady=5)
        self.tc_exp_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.tc_exp_var, width=40).grid(row=4, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Skutočný výsledok:").grid(row=5, column=0, sticky="w", padx=5, pady=5)
        self.tc_act_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.tc_act_var, width=40).grid(row=5, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Stav:").grid(row=6, column=0, sticky="w", padx=5, pady=5)
        self.tc_status_var = tk.StringVar(value="NOT RUN")
        self.tc_status_combo = ttk.Combobox(
            form_frame,
            textvariable=self.tc_status_var,
            values=["PASSED", "FAILED", "NOT RUN"],
            state="readonly",
            width=37,
        )
        self.tc_status_combo.grid(row=6, column=1, padx=5, pady=5)

        ttk.Button(form_frame, text="Uložiť nový TC", command=self.add_tc).grid(
            row=7, column=1, sticky="e", padx=5, pady=5
        )
        ttk.Button(form_frame, text="Uložiť zmeny TC", command=self.update_tc).grid(
            row=8, column=1, sticky="e", padx=5, pady=5
        )
        ttk.Button(form_frame, text="Vymazať tento TC", command=self.delete_selected_tc).grid(
            row=9, column=1, sticky="e", padx=5, pady=5
        )

        list_frame = ttk.LabelFrame(self.tc_tab, text="Zoznam Test Cases")
        list_frame.pack(side="top", fill="both", expand=True, padx=10, pady=5)

        filter_frame = ttk.Frame(list_frame)
        filter_frame.pack(side="top", fill="x", padx=5, pady=5)

        ttk.Label(filter_frame, text="Filter podľa statusu:").pack(side="left")
        self.tc_filter_var = tk.StringVar(value="ALL")
        self.tc_filter_combo = ttk.Combobox(
            filter_frame,
            textvariable=self.tc_filter_var,
            values=["ALL", "PASSED", "FAILED", "NOT RUN"],
            state="readonly",
            width=10,
        )
        self.tc_filter_combo.pack(side="left", padx=5)
        ttk.Button(filter_frame, text="Použiť filter", command=self.refresh_tc_list).pack(side="left")

        middle = ttk.Frame(list_frame)
        middle.pack(side="top", fill="both", expand=True, padx=5, pady=5)

        self.tc_list = tk.Listbox(middle, width=50)
        self.tc_list.pack(side="left", fill="both", expand=True, padx=(0, 5), pady=5)
        self.tc_list.bind("<<ListboxSelect>>", self.show_tc_detail)

        self.tc_detail = tk.Text(middle, width=50)
        self.tc_detail.pack(side="left", fill="both", expand=True, padx=(5, 0), pady=5)

        self._styled_text_widgets.append(self.tc_steps_text)
        self._styled_text_widgets.append(self.tc_detail)
        self._styled_listbox_widgets.append(self.tc_list)

        self.refresh_tc_ts_combobox()
        self.refresh_tc_list()

    def refresh_tc_ts_combobox(self):
        ids = [ts["id"] for ts in self.data["test_scenarios"]]
        self.tc_ts_combo["values"] = ids

    def add_tc(self):
        title = self.tc_title_var.get().strip()
        pre = self.tc_pre_var.get().strip()
        ts_id = self.tc_ts_var.get().strip() or None
        steps_raw = self.tc_steps_text.get("1.0", "end").strip()
        expected = self.tc_exp_var.get().strip()
        actual = self.tc_act_var.get().strip()
        status = self.tc_status_var.get().strip() or "NOT RUN"

        if not title:
            messagebox.showerror("Chyba", "Názov TC nemôže byť prázdny.")
            return
        if not steps_raw:
            messagebox.showerror("Chyba", "Musíš zadať aspoň jeden krok.")
            return

        steps = [s.strip() for s in steps_raw.splitlines() if s.strip()]

        tc_id = generate_id("TC", self.data["test_cases"])
        tc = {
            "id": tc_id,
            "title": title,
            "preconditions": pre,
            "ts_id": ts_id,
            "steps": steps,
            "expected": expected,
            "actual": actual,
            "status": status.upper(),
        }

        self.data["test_cases"].append(tc)
        save_data(self.data)

        self.refresh_bug_tc_combobox()

        self.tc_title_var.set("")
        self.tc_pre_var.set("")
        self.tc_ts_var.set("")
        self.tc_steps_text.delete("1.0", "end")
        self.tc_exp_var.set("")
        self.tc_act_var.set("")
        self.tc_status_var.set("NOT RUN")
        self.selected_tc_id = None

        self.refresh_tc_list()
        self.apply_theme()
        messagebox.showinfo("OK", f"Test case {tc_id} uložený.")

    def refresh_tc_list(self, event=None):
        self.tc_list.delete(0, "end")
        self.tc_detail.delete("1.0", "end")

        filter_val = self.tc_filter_var.get()
        for tc in self.data["test_cases"]:
            if filter_val != "ALL" and tc["status"] != filter_val:
                continue
            line = f"{tc['id']} – {tc['title']} [{tc['status']}]"
            self.tc_list.insert("end", line)

        self.selected_tc_id = None

    def show_tc_detail(self, event):
        selection = self.tc_list.curselection()
        if not selection:
            return
        index = selection[0]

        filter_val = self.tc_filter_var.get()
        filtered = []
        for tc in self.data["test_cases"]:
            if filter_val != "ALL" and tc["status"] != filter_val:
                continue
            filtered.append(tc)
        if index >= len(filtered):
            return
        tc = filtered[index]

        self.selected_tc_id = tc["id"]

        self.tc_detail.delete("1.0", "end")
        lines = []
        lines.append(f"ID: {tc['id']}")
        lines.append(f"Názov: {tc['title']}")
        lines.append(f"TS: {tc['ts_id']}")
        lines.append(f"Predpoklady: {tc['preconditions']}")
        lines.append("")
        lines.append("Kroky:")
        for i, step in enumerate(tc["steps"], start=1):
            lines.append(f"  {i}. {step}")
        lines.append("")
        lines.append(f"Očakávaný výsledok: {tc['expected']}")
        lines.append(f"Skutočný výsledok: {tc['actual']}")
        lines.append(f"Stav: {tc['status']}")

        self.tc_detail.insert("1.0", "\n".join(lines))

        # naplň formulár
        self.tc_title_var.set(tc["title"])
        self.tc_pre_var.set(tc["preconditions"])
        self.tc_ts_var.set(tc["ts_id"] or "")
        self.tc_steps_text.delete("1.0", "end")
        self.tc_steps_text.insert("1.0", "\n".join(tc["steps"]))
        self.tc_exp_var.set(tc["expected"])
        self.tc_act_var.set(tc["actual"])
        self.tc_status_var.set(tc["status"])

    def update_tc(self):
        if not self.selected_tc_id:
            messagebox.showerror("Chyba", "Najprv vyber test case zo zoznamu.")
            return

        title = self.tc_title_var.get().strip()
        pre = self.tc_pre_var.get().strip()
        ts_id = self.tc_ts_var.get().strip() or None
        steps_raw = self.tc_steps_text.get("1.0", "end").strip()
        expected = self.tc_exp_var.get().strip()
        actual = self.tc_act_var.get().strip()
        status = self.tc_status_var.get().strip() or "NOT RUN"

        if not title:
            messagebox.showerror("Chyba", "Názov TC nemôže byť prázdny.")
            return
        if not steps_raw:
            messagebox.showerror("Chyba", "Musíš zadať aspoň jeden krok.")
            return

        steps = [s.strip() for s in steps_raw.splitlines() if s.strip()]

        for tc in self.data["test_cases"]:
            if tc["id"] == self.selected_tc_id:
                tc["title"] = title
                tc["preconditions"] = pre
                tc["ts_id"] = ts_id
                tc["steps"] = steps
                tc["expected"] = expected
                tc["actual"] = actual
                tc["status"] = status.upper()
                break

        save_data(self.data)
        self.refresh_tc_list()
        self.refresh_bug_tc_combobox()
        self.apply_theme()
        messagebox.showinfo("OK", f"Test case {self.selected_tc_id} bol upravený.")

    def delete_selected_tc(self):
        if not self.selected_tc_id:
            messagebox.showerror("Chyba", "Najprv vyber test case zo zoznamu.")
            return

        tc_to_delete = None
        for tc in self.data["test_cases"]:
            if tc["id"] == self.selected_tc_id:
                tc_to_delete = tc
                break

        if not tc_to_delete:
            messagebox.showerror("Chyba", "Test case sa nenašiel.")
            return

        if not messagebox.askyesno(
            "Vymazať TC",
            f"Naozaj chceš vymazať {tc_to_delete['id']} – {tc_to_delete['title']}?\n\nBugy, ktoré naň odkazujú, zostanú s rovnakým ID TC."
        ):
            return

        self.data["test_cases"] = [tc for tc in self.data["test_cases"] if tc["id"] != self.selected_tc_id]
        save_data(self.data)

        self.tc_title_var.set("")
        self.tc_pre_var.set("")
        self.tc_ts_var.set("")
        self.tc_steps_text.delete("1.0", "end")
        self.tc_exp_var.set("")
        self.tc_act_var.set("")
        self.tc_status_var.set("NOT RUN")
        self.selected_tc_id = None

        self.refresh_tc_list()
        self.refresh_bug_tc_combobox()
        messagebox.showinfo("OK", "Test case bol vymazaný.")

    # ===== TAB – BUG REPORTY =====
    def create_bug_tab(self):
        self.bug_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.bug_tab, text="Bug Reports")

        form_frame = ttk.LabelFrame(self.bug_tab, text="Pridať / Upraviť Bug Report")
        form_frame.pack(side="top", fill="x", padx=10, pady=10)

        ttk.Label(form_frame, text="Názov chyby:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.bug_title_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.bug_title_var, width=50).grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="ID Test Case (TC):").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        self.bug_tc_var = tk.StringVar()
        self.bug_tc_combo = ttk.Combobox(form_frame, textvariable=self.bug_tc_var, width=47, state="readonly")
        self.bug_tc_combo.grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Kroky k reprodukcii:").grid(row=2, column=0, sticky="nw", padx=5, pady=5)
        self.bug_steps_text = tk.Text(form_frame, height=4, width=50)
        self.bug_steps_text.grid(row=2, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Očakávaný výsledok:").grid(row=3, column=0, sticky="w", padx=5, pady=5)
        self.bug_exp_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.bug_exp_var, width=50).grid(row=3, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Skutočný výsledok:").grid(row=4, column=0, sticky="w", padx=5, pady=5)
        self.bug_act_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.bug_act_var, width=50).grid(row=4, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Severity:").grid(row=5, column=0, sticky="w", padx=5, pady=5)
        self.bug_sev_var = tk.StringVar(value="Medium")
        self.bug_sev_combo = ttk.Combobox(
            form_frame,
            textvariable=self.bug_sev_var,
            values=["Low", "Medium", "High", "Critical"],
            state="readonly",
            width=47,
        )
        self.bug_sev_combo.grid(row=5, column=1, padx=5, pady=5)

        ttk.Label(form_frame, text="Poznámka:").grid(row=6, column=0, sticky="w", padx=5, pady=5)
        self.bug_note_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.bug_note_var, width=50).grid(row=6, column=1, padx=5, pady=5)

        # Screenshot – nový riadok
        ttk.Label(form_frame, text="Screenshot:").grid(row=7, column=0, sticky="w", padx=5, pady=5)
        self.bug_screenshot_label = ttk.Label(form_frame, text="Žiadny súbor nevybraný")
        self.bug_screenshot_label.grid(row=7, column=1, sticky="w", padx=5, pady=5)

        ttk.Button(form_frame, text="Vybrať screenshot", command=self.select_bug_screenshot).grid(
            row=8, column=0, sticky="w", padx=5, pady=5
        )

        ttk.Button(form_frame, text="Uložiť Bug", command=self.add_bug).grid(
            row=9, column=1, sticky="e", padx=5, pady=5
        )
        ttk.Button(form_frame, text="Uložiť zmeny Bug", command=self.update_bug).grid(
            row=10, column=1, sticky="e", padx=5, pady=5
        )
        ttk.Button(form_frame, text="Vymazať tento Bug", command=self.delete_selected_bug).grid(
            row=11, column=1, sticky="e", padx=5, pady=5
        )

        list_frame = ttk.LabelFrame(self.bug_tab, text="Zoznam Bug Reportov")
        list_frame.pack(side="top", fill="both", expand=True, padx=10, pady=10)

        middle = ttk.Frame(list_frame)
        middle.pack(side="top", fill="both", expand=True, padx=5, pady=5)

        self.bug_list = tk.Listbox(middle, width=50)
        self.bug_list.pack(side="left", fill="both", expand=True, padx=(0, 5), pady=5)
        self.bug_list.bind("<<ListboxSelect>>", self.show_bug_detail)

        self.bug_detail = tk.Text(middle, width=50)
        self.bug_detail.pack(side="left", fill="both", expand=True, padx=(5, 0), pady=5)

        self._styled_text_widgets.append(self.bug_steps_text)
        self._styled_text_widgets.append(self.bug_detail)
        self._styled_listbox_widgets.append(self.bug_list)

        self.refresh_bug_tc_combobox()
        self.refresh_bug_list()

    def select_bug_screenshot(self):
        file_path = filedialog.askopenfilename(
            title="Vyber screenshot",
            filetypes=[
                ("Image files", "*.png *.jpg *.jpeg *.gif *.bmp"),
                ("All files", "*.*"),
            ]
        )
        if file_path:
            self.bug_screenshot_path = file_path
            file_name = os.path.basename(file_path)
            self.bug_screenshot_label.config(text=f"Vybraný: {file_name}")

    def refresh_bug_tc_combobox(self):
        ids = [tc["id"] for tc in self.data["test_cases"]]
        self.bug_tc_combo["values"] = ids

    def add_bug(self):
        title = self.bug_title_var.get().strip()
        tc_id = self.bug_tc_var.get().strip()
        steps_raw = self.bug_steps_text.get("1.0", "end").strip()
        expected = self.bug_exp_var.get().strip()
        actual = self.bug_act_var.get().strip()
        severity = self.bug_sev_var.get().strip()
        note = self.bug_note_var.get().strip()
        screenshot = self.bug_screenshot_path

        if not title:
            messagebox.showerror("Chyba", "Názov chyby nemôže byť prázdny.")
            return
        if not steps_raw:
            messagebox.showerror("Chyba", "Musíš zadať kroky k reprodukcii.")
            return

        steps = [s.strip() for s in steps_raw.splitlines() if s.strip()]

        bug_id = generate_id("BUG", self.data["bug_reports"])
        bug = {
            "id": bug_id,
            "title": title,
            "related_tc": tc_id,
            "steps": steps,
            "expected": expected,
            "actual": actual,
            "severity": severity,
            "note": note,
            "screenshot": screenshot,
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }

        self.data["bug_reports"].append(bug)

        if tc_id:
            for tc in self.data["test_cases"]:
                if tc["id"] == tc_id:
                    tc["status"] = "FAILED"
                    break

        save_data(self.data)

        self.bug_title_var.set("")
        self.bug_tc_var.set("")
        self.bug_steps_text.delete("1.0", "end")
        self.bug_exp_var.set("")
        self.bug_act_var.set("")
        self.bug_sev_var.set("Medium")
        self.bug_note_var.set("")
        self.bug_screenshot_path = None
        self.bug_screenshot_label.config(text="Žiadny súbor nevybraný")
        self.selected_bug_id = None

        self.refresh_bug_list()
        self.refresh_tc_list()
        self.apply_theme()
        messagebox.showinfo("OK", f"Bug {bug_id} uložený.")

    def refresh_bug_list(self):
        self.bug_list.delete(0, "end")
        self.bug_detail.delete("1.0", "end")
        for bug in self.data["bug_reports"]:
            line = f"{bug['id']} – {bug['title']} [{bug['severity']}]"
            self.bug_list.insert("end", line)
        self.selected_bug_id = None

    def show_bug_detail(self, event):
        selection = self.bug_list.curselection()
        if not selection:
            return
        index = selection[0]
        if index >= len(self.data["bug_reports"]):
            return
        bug = self.data["bug_reports"][index]
        self.selected_bug_id = bug["id"]

        self.bug_detail.delete("1.0", "end")
        lines = []
        lines.append(f"ID: {bug['id']}")
        lines.append(f"Názov: {bug['title']}")
        lines.append(f"Test Case: {bug['related_tc']}")
        lines.append(f"Severity: {bug['severity']}")
        lines.append(f"Vytvorené: {bug['created_at']}")
        lines.append("")
        lines.append("Kroky k reprodukcii:")
        for i, step in enumerate(bug["steps"], start=1):
            lines.append(f"  {i}. {step}")
        lines.append("")
        lines.append(f"Očakávaný výsledok: {bug['expected']}")
        lines.append(f"Skutočný výsledok: {bug['actual']}")
        screenshot = bug.get("screenshot")
        if screenshot:
            lines.append("")
            lines.append(f"Screenshot: {screenshot}")
        if bug["note"]:
            lines.append("")
            lines.append(f"Poznámka: {bug['note']}")

        self.bug_detail.insert("1.0", "\n".join(lines))

        # naplň formulár
        self.bug_title_var.set(bug["title"])
        self.bug_tc_var.set(bug["related_tc"])
        self.bug_steps_text.delete("1.0", "end")
        self.bug_steps_text.insert("1.0", "\n".join(bug["steps"]))
        self.bug_exp_var.set(bug["expected"])
        self.bug_act_var.set(bug["actual"])
        self.bug_sev_var.set(bug["severity"])
        self.bug_note_var.set(bug["note"])
        self.bug_screenshot_path = bug.get("screenshot")
        if self.bug_screenshot_path:
            file_name = os.path.basename(self.bug_screenshot_path)
            self.bug_screenshot_label.config(text=f"Vybraný: {file_name}")
        else:
            self.bug_screenshot_label.config(text="Žiadny súbor nevybraný")

    def update_bug(self):
        if not self.selected_bug_id:
            messagebox.showerror("Chyba", "Najprv vyber bug zo zoznamu.")
            return

        title = self.bug_title_var.get().strip()
        tc_id = self.bug_tc_var.get().strip()
        steps_raw = self.bug_steps_text.get("1.0", "end").strip()
        expected = self.bug_exp_var.get().strip()
        actual = self.bug_act_var.get().strip()
        severity = self.bug_sev_var.get().strip()
        note = self.bug_note_var.get().strip()
        screenshot = self.bug_screenshot_path

        if not title:
            messagebox.showerror("Chyba", "Názov chyby nemôže byť prázdny.")
            return
        if not steps_raw:
            messagebox.showerror("Chyba", "Musíš zadať kroky k reprodukcii.")
            return

        steps = [s.strip() for s in steps_raw.splitlines() if s.strip()]

        for bug in self.data["bug_reports"]:
            if bug["id"] == self.selected_bug_id:
                bug["title"] = title
                bug["related_tc"] = tc_id
                bug["steps"] = steps
                bug["expected"] = expected
                bug["actual"] = actual
                bug["severity"] = severity
                bug["note"] = note
                bug["screenshot"] = screenshot
                break

        if tc_id:
            for tc in self.data["test_cases"]:
                if tc["id"] == tc_id:
                    tc["status"] = "FAILED"
                    break

        save_data(self.data)
        self.refresh_bug_list()
        self.refresh_tc_list()
        self.apply_theme()
        messagebox.showinfo("OK", f"Bug {self.selected_bug_id} bol upravený.")

    def delete_selected_bug(self):
        if not self.selected_bug_id:
            messagebox.showerror("Chyba", "Najprv vyber bug zo zoznamu.")
            return

        bug_to_delete = None
        for bug in self.data["bug_reports"]:
            if bug["id"] == self.selected_bug_id:
                bug_to_delete = bug
                break

        if not bug_to_delete:
            messagebox.showerror("Chyba", "Bug sa nenašiel.")
            return

        if not messagebox.askyesno(
            "Vymazať Bug",
            f"Naozaj chceš vymazať {bug_to_delete['id']} – {bug_to_delete['title']}?"
        ):
            return

        self.data["bug_reports"] = [bug for bug in self.data["bug_reports"] if bug["id"] != self.selected_bug_id]
        save_data(self.data)

        self.bug_title_var.set("")
        self.bug_tc_var.set("")
        self.bug_steps_text.delete("1.0", "end")
        self.bug_exp_var.set("")
        self.bug_act_var.set("")
        self.bug_sev_var.set("Medium")
        self.bug_note_var.set("")
        self.bug_screenshot_path = None
        self.bug_screenshot_label.config(text="Žiadny súbor nevybraný")
        self.selected_bug_id = None

        self.refresh_bug_list()
        messagebox.showinfo("OK", "Bug bol vymazaný.")

    # ===== EXPORTY =====
    def export_to_txt(self):
        filename = "qa_export.txt"
        with open(filename, "w", encoding="utf-8") as f:
            f.write("=== TEST SCENÁRE ===\n\n")
            for ts in self.data["test_scenarios"]:
                f.write(f"{ts['id']} – {ts['title']}\n")
                if ts["description"]:
                    f.write(f"Popis: {ts['description']}\n")
                f.write("-" * 60 + "\n")

            f.write("\n=== TEST CASES ===\n\n")
            for tc in self.data["test_cases"]:
                f.write(f"ID: {tc['id']}\n")
                f.write(f"Názov: {tc['title']}\n")
                f.write(f"TS: {tc['ts_id']}\n")
                f.write(f"Predpoklady: {tc['preconditions']}\n")
                f.write("Kroky:\n")
                for i, step in enumerate(tc["steps"], start=1):
                    f.write(f"  {i}. {step}\n")
                f.write(f"Očakávaný výsledok: {tc['expected']}\n")
                f.write(f"Skutočný výsledok: {tc['actual']}\n")
                f.write(f"Stav: {tc['status']}\n")
                f.write("-" * 60 + "\n")

            f.write("\n=== BUG REPORTS ===\n\n")
            for bug in self.data["bug_reports"]:
                f.write(f"ID: {bug['id']}\n")
                f.write(f"Názov: {bug['title']}\n")
                f.write(f"Test Case: {bug['related_tc']}\n")
                f.write(f"Severity: {bug['severity']}\n")
                f.write("Kroky k reprodukcii:\n")
                for i, step in enumerate(bug["steps"], start=1):
                    f.write(f"  {i}. {step}\n")
                f.write(f"Očakávaný výsledok: {bug['expected']}\n")
                f.write(f"Skutočný výsledok: {bug['actual']}\n")
                screenshot = bug.get("screenshot")
                if screenshot:
                    f.write(f"Screenshot: {screenshot}\n")
                if bug["note"]:
                    f.write(f"Poznámka: {bug['note']}\n")
                f.write(f"Vytvorené: {bug['created_at']}\n")
                f.write("-" * 60 + "\n")

        messagebox.showinfo("Export", f"TXT export vytvorený: {filename}")

    def export_to_html(self):
        filename = "qa_export.html"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(
                "<html><head><meta charset='utf-8'>"
                "<title>QA Export</title>"
                "<style>"
                "body{font-family:Arial, sans-serif;}"
                "h1,h2{color:#333;}"
                ".section{margin-bottom:30px;}"
                ".card{border:1px solid #ccc;padding:10px;margin:5px 0;}"
                ".bug{border-color:#f00;}"
                "table{border-collapse:collapse;width:100%;margin:5px 0;}"
                "th,td{border:1px solid #ccc;padding:4px;font-size:12px;}"
                "th{background:#f5f5f5;}"
                "</style>"
                "</head><body>"
            )
            f.write("<h1>QA Export</h1>")
            f.write(f"<p>Vygenerované: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>")

            # Test scenáre
            f.write("<div class='section'><h2>Test Scenáre</h2>")
            for ts in self.data["test_scenarios"]:
                f.write("<div class='card'>")
                f.write(f"<strong>{ts['id']}</strong> – {ts['title']}<br>")
                if ts["description"]:
                    f.write(f"<em>Popis:</em> {ts['description']}<br>")
                f.write("</div>")
            f.write("</div>")

            # Test Cases
            f.write("<div class='section'><h2>Test Cases</h2>")
            for tc in self.data["test_cases"]:
                f.write("<div class='card'>")
                f.write(f"<strong>{tc['id']}</strong> – {tc['title']}<br>")
                f.write(f"<strong>TS:</strong> {tc['ts_id']}<br>")
                f.write(f"<strong>Predpoklady:</strong> {tc['preconditions']}<br>")
                f.write("<strong>Kroky:</strong><ol>")
                for step in tc["steps"]:
                    f.write(f"<li>{step}</li>")
                f.write("</ol>")
                f.write(f"<strong>Očakávaný:</strong> {tc['expected']}<br>")
                f.write(f"<strong>Skutočný:</strong> {tc['actual']}<br>")
                f.write(f"<strong>Stav:</strong> {tc['status']}<br>")
                f.write("</div>")
            f.write("</div>")

            # Bug Reports
            f.write("<div class='section'><h2>Bug Reports</h2>")
            for bug in self.data["bug_reports"]:
                f.write("<div class='card bug'>")
                f.write(f"<strong>{bug['id']}</strong> – {bug['title']}<br>")
                f.write(f"<strong>Test Case:</strong> {bug['related_tc']}<br>")
                f.write(f"<strong>Severity:</strong> {bug['severity']}<br>")
                f.write(f"<strong>Vytvorené:</strong> {bug['created_at']}<br>")
                f.write("<strong>Kroky k reprodukcii:</strong><ol>")
                for step in bug["steps"]:
                    f.write(f"<li>{step}</li>")
                f.write("</ol>")
                f.write(f"<strong>Očakávaný:</strong> {bug['expected']}<br>")
                f.write(f"<strong>Skutočný:</strong> {bug['actual']}<br>")

                screenshot = bug.get("screenshot")
                if screenshot:
                    web_path = screenshot.replace("\\", "/")
                    f.write("<strong>Screenshot:</strong><br>")
                    f.write(
                        f"<img src='{web_path}' "
                        f"style='max-width:400px; max-height:300px; border:1px solid #ccc;'><br>"
                    )

                if bug["note"]:
                    f.write(f"<strong>Poznámka:</strong> {bug['note']}<br>")
                f.write("</div>")
            f.write("</div>")

            f.write("</body></html>")

        messagebox.showinfo("Export", f"HTML export vytvorený: {filename}")

    def export_to_word(self):
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            messagebox.showerror(
                "Chýbajúci balík",
                "Na export do Wordu potrebuješ balík 'python-docx'.\n"
                "Nainštaluj ho napríklad príkazom:\n\npip install python-docx"
            )
            return

        doc = Document()

        # Hlavný nadpis
        doc.add_heading("QA Test Report", level=1)
        doc.add_paragraph(f"Vygenerované: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # ===== TEST SCENÁRE =====
        doc.add_page_break()
        doc.add_heading("Test Scenáre", level=2)

        if self.data["test_scenarios"]:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "ID"
            hdr_cells[1].text = "Názov"
            hdr_cells[2].text = "Popis"

            for ts in self.data["test_scenarios"]:
                row_cells = table.add_row().cells
                row_cells[0].text = ts["id"]
                row_cells[1].text = ts["title"] or ""
                row_cells[2].text = ts["description"] or ""
        else:
            doc.add_paragraph("Žiadne test scenáre.", style="Intense Quote")

        # ===== TEST CASES =====
        doc.add_page_break()
        doc.add_heading("Test Cases", level=2)

        if self.data["test_cases"]:
            table = doc.add_table(rows=1, cols=8)
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "ID"
            hdr_cells[1].text = "Názov"
            hdr_cells[2].text = "TS"
            hdr_cells[3].text = "Predpoklady"
            hdr_cells[4].text = "Kroky"
            hdr_cells[5].text = "Očakávaný výsledok"
            hdr_cells[6].text = "Skutočný výsledok"
            hdr_cells[7].text = "Stav"

            for tc in self.data["test_cases"]:
                row_cells = table.add_row().cells
                row_cells[0].text = tc["id"]
                row_cells[1].text = tc["title"] or ""
                row_cells[2].text = tc["ts_id"] or ""
                row_cells[3].text = tc["preconditions"] or ""
                steps_text = "\n".join(tc["steps"]) if tc["steps"] else ""
                row_cells[4].text = steps_text
                row_cells[5].text = tc["expected"] or ""
                row_cells[6].text = tc["actual"] or ""
                row_cells[7].text = tc["status"] or ""
        else:
            doc.add_paragraph("Žiadne Test Cases.", style="Intense Quote")

        # ===== BUG REPORTS =====
        doc.add_page_break()
        doc.add_heading("Bug Reports", level=2)

        if self.data["bug_reports"]:
            table = doc.add_table(rows=1, cols=9)
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "ID"
            hdr_cells[1].text = "Názov"
            hdr_cells[2].text = "Test Case"
            hdr_cells[3].text = "Severity"
            hdr_cells[4].text = "Kroky k reprodukcii"
            hdr_cells[5].text = "Očakávaný výsledok"
            hdr_cells[6].text = "Skutočný výsledok"
            hdr_cells[7].text = "Poznámka"
            hdr_cells[8].text = "Vytvorené"

            for bug in self.data["bug_reports"]:
                row_cells = table.add_row().cells
                row_cells[0].text = bug["id"]
                row_cells[1].text = bug["title"] or ""
                row_cells[2].text = bug["related_tc"] or ""
                row_cells[3].text = bug["severity"] or ""
                steps_text = "\n".join(bug["steps"]) if bug["steps"] else ""
                row_cells[4].text = steps_text
                row_cells[5].text = bug["expected"] or ""
                row_cells[6].text = bug["actual"] or ""
                row_cells[7].text = bug["note"] or ""
                row_cells[8].text = bug["created_at"] or ""

                screenshot = bug.get("screenshot")
                if screenshot:
                    p = doc.add_paragraph()
                    p.add_run(f"Screenshot pre {bug['id']} – {bug['title']}:\n")
                    try:
                        doc.add_picture(screenshot, width=Inches(3))
                    except Exception as e:
                        p.add_run(f"(Nepodarilo sa vložiť obrázok: {e})")
                    doc.add_paragraph("")
        else:
            doc.add_paragraph("Žiadne bug reporty.", style="Intense Quote")

        filename = "qa_export_professional.docx"
        doc.save(filename)
        messagebox.showinfo("Export", f"Word export (profi) vytvorený: {filename}")

    def export_to_pdf(self):
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import cm
            from reportlab.lib.utils import ImageReader
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            messagebox.showerror(
                "Chýbajúci balík",
                "Na export do PDF potrebuješ balík 'reportlab'.\n"
                "Nainštaluj ho napríklad príkazom:\n\npip install reportlab"
            )
            return

        # Registrujeme Arial z arial.ttf v tom istom priečinku
        try:
            pdfmetrics.registerFont(TTFont("Arial", "arial.ttf"))
            base_font_name = "Arial"
        except Exception as e:
            messagebox.showwarning(
                "Font",
                f"Nepodarilo sa načítať arial.ttf, použije sa predvolený font (bez diakritiky).\n\n{e}"
            )
            base_font_name = "Helvetica"

        filename = "qa_export.pdf"
        c = canvas.Canvas(filename, pagesize=A4)
        width, height = A4
        y = height - 2 * cm

        def new_page():
            nonlocal y
            c.showPage()
            y = height - 2 * cm
            c.setFont(base_font_name, 10)

        def write_line(text="", size=10):
            nonlocal y
            if y < 2 * cm:
                new_page()
            c.setFont(base_font_name, size)
            c.drawString(2 * cm, y, text)
            y -= 0.6 * cm

        c.setFont(base_font_name, 10)

        # Nadpis
        write_line("QA Test Report", size=16)
        write_line(f"Vygenerované: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", size=10)
        y -= 0.5 * cm

        # Test scenáre
        write_line("Test Scenáre", size=14)
        y -= 0.2 * cm
        for ts in self.data["test_scenarios"]:
            write_line(f"{ts['id']} – {ts['title']}", size=11)
            if ts["description"]:
                for line in ts["description"].splitlines():
                    write_line(f"  {line}", size=10)
            y -= 0.2 * cm
        y -= 0.5 * cm

        # Test Cases
        write_line("Test Cases", size=14)
        y -= 0.2 * cm
        for tc in self.data["test_cases"]:
            write_line(f"{tc['id']} – {tc['title']}", size=11)
            write_line(f"TS: {tc['ts_id']}", size=10)
            write_line(f"Predpoklady: {tc['preconditions']}", size=10)
            write_line("Kroky:", size=10)
            for i, step in enumerate(tc["steps"], start=1):
                write_line(f"  {i}. {step}", size=10)
            write_line(f"Očakávaný výsledok: {tc['expected']}", size=10)
            write_line(f"Skutočný výsledok: {tc['actual']}", size=10)
            write_line(f"Stav: {tc['status']}", size=10)
            y -= 0.4 * cm
        y -= 0.5 * cm

        # Bug Reports
        write_line("Bug Reports", size=14)
        y -= 0.2 * cm
        for bug in self.data["bug_reports"]:
            write_line(f"{bug['id']} – {bug['title']}", size=11)
            write_line(f"Test Case: {bug['related_tc']}", size=10)
            write_line(f"Severity: {bug['severity']}", size=10)
            write_line(f"Vytvorené: {bug['created_at']}", size=10)
            write_line("Kroky k reprodukcii:", size=10)
            for i, step in enumerate(bug["steps"], start=1):
                write_line(f"  {i}. {step}", size=10)
            write_line(f"Očakávaný výsledok: {bug['expected']}", size=10)
            write_line(f"Skutočný výsledok: {bug['actual']}", size=10)

            screenshot = bug.get("screenshot")
            if screenshot:
                y -= 0.2 * cm
                write_line("Screenshot:", size=10)
                if y < 8 * cm:
                    new_page()
                try:
                    img = ImageReader(screenshot)
                    iw, ih = img.getSize()
                    max_w = width - 4 * cm
                    max_h = 8 * cm
                    scale = min(max_w / iw, max_h / ih, 1.0)
                    img_w = iw * scale
                    img_h = ih * scale
                    if y - img_h < 2 * cm:
                        new_page()
                    c.drawImage(img, 2 * cm, y - img_h, width=img_w, height=img_h)
                    y -= img_h + 0.5 * cm
                except Exception as e:
                    write_line(f"(Nepodarilo sa vložiť obrázok: {e})", size=9)
            if bug["note"]:
                write_line(f"Poznámka: {bug['note']}", size=10)
            y -= 0.4 * cm

        c.save()
        messagebox.showinfo("Export", f"PDF export vytvorený: {filename}")

    # ===== RESET DATABÁZY =====
    def reset_database(self):
        if not messagebox.askyesno(
            "Reset databázy",
            "Naozaj chceš vymazať všetky uložené dáta?\n"
            "Súbor qa_data_gui.json bude zmazaný a všetko sa vynuluje."
        ):
            return

        if os.path.exists(DATA_FILE):
            try:
                os.remove(DATA_FILE)
            except OSError as e:
                messagebox.showerror("Chyba", f"Nepodarilo sa zmazať súbor: {e}")
                return

        self.data = {"test_scenarios": [], "test_cases": [], "bug_reports": []}

        self.refresh_ts_list()
        self.refresh_tc_ts_combobox()
        self.refresh_tc_list()
        self.refresh_bug_tc_combobox()
        self.refresh_bug_list()
        self.selected_ts_index = None
        self.selected_tc_id = None
        self.selected_bug_id = None
        self.bug_screenshot_path = None
        if hasattr(self, "bug_screenshot_label"):
            self.bug_screenshot_label.config(text="Žiadny súbor nevybraný")
        messagebox.showinfo("Reset", "Databáza bola vymazaná.")


if __name__ == "__main__":
    app = QAApp()
    app.mainloop()
